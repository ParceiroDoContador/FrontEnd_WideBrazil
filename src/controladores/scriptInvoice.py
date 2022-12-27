import boto3
import requests
import json
from datetime import datetime, timedelta, date
import unidecode
import shutil
import base64
import random
from docx.oxml.shared import OxmlElement
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches, RGBColor
from docx2pdf import convert

#=================== Verificação de Liberão ========================#
liberacao = requests.get("https://gliciojunior.notion.site/WIDE-5ed9ee76906a444187fccaaba35702de")
print(f"liberacao: {liberacao}")
if str(liberacao) == "<Response [200]>":
    #======================= Categorias ==============================#
    categoria_receber_salario = "1.01.97"
    categoria_receber_comissao = "1.01.02"
    categoria_receber_dsr = "1.01.03"
    categoria_receber_alimentacao = "1.02.01"
    categoria_receber_reembolso_despesas = "1.04.02"
    categoria_receber_reembolso_saude = "2.01.01"
    categoria_receber_inss = "1.01.96"
    categoria_receber_adiantamento = "1.03.03"
    categoria_receber_irrf = ""
    categoria_receber_previdencia = "1.03.26"
    categoria_receber_inss_empresa = "1.01.99"
    categoria_receber_fgts = "1.01.98"
    categoria_receber_liquido = "1.04.03"
    categoria_receber_ferias = "1.01.94"
    categoria_receber_seguro = "1.04.06"
    categoria_receber_decimo = "1.01.90"
    categoria_receber_flash = "1.01.91"

    categoria_invoice = "1.01.89"

    app_key = '3068480598183'
    app_secret = '91ed53d6746eb516fd6239186c82ad65'
    def buscar_codigo_cliente_teste(nome):
        nome = unidecode.unidecode(nome).upper()
        pagina = 1
        total_de_paginas = 1
        while pagina <= total_de_paginas:
            url = "https://app.omie.com.br/api/v1/geral/clientes/"
            payload = json.dumps({
                                    "call": "ListarClientes",
                                    "app_key": app_key,
                                    "app_secret": app_secret,
                                    "param":[
                                                {
                                                    "pagina": pagina,
                                                    "registros_por_pagina": 500,
                                                    "apenas_importado_api": "N"
                                                }
                                            ]
                                })
            headers ={
                        'Content-Type': 'application/json'
                    }
            response = requests.request("POST", url, headers=headers, data=payload)
            response = response.json()
            pagina = response["pagina"]
            total_de_paginas = response["total_de_paginas"]
            clientes_cadastro = response["clientes_cadastro"]
            for cliente in clientes_cadastro:
                try:
                    contato = cliente["contato"]
                    contato = unidecode.unidecode(contato).upper()
                    if contato == nome:
                        codigo_cliente_omie = cliente["codigo_cliente_omie"]
                        break
                except:
                    pass        
            pagina += 1
        return codigo_cliente_omie
    def buscar_codigo_cliente(nome):
        nome = unidecode.unidecode(nome).upper()
        pagina = 1
        total_de_paginas = 1
        while pagina <= total_de_paginas:
            url = "https://app.omie.com.br/api/v1/geral/clientes/"
            payload = json.dumps({
                                    "call": "ListarClientes",
                                    "app_key": app_key,
                                    "app_secret": app_secret,
                                    "param":[
                                                {
                                                    "pagina": pagina,
                                                    "registros_por_pagina": 500,
                                                    "apenas_importado_api": "N"
                                                }
                                            ]
                                })
            headers ={
                        'Content-Type': 'application/json'
                    }
            response = requests.request("POST", url, headers=headers, data=payload)
            response = response.json()
            pagina = response["pagina"]
            total_de_paginas = response["total_de_paginas"]
            clientes_cadastro = response["clientes_cadastro"]
            for cliente in clientes_cadastro:  
                razao_social = cliente["razao_social"]
                nome_fantasia = cliente["nome_fantasia"]
                razao_social = unidecode.unidecode(razao_social).upper()
                nome_fantasia = unidecode.unidecode(nome_fantasia).upper()  
                print(f'nome: {nome} - razao_social: {razao_social} - nome_fantasia: {nome_fantasia}')              
                if razao_social == nome or nome_fantasia == nome:
                    codigo_cliente_omie = cliente["codigo_cliente_omie"]
                    break
            pagina += 1
        return codigo_cliente_omie
    def gerar_invoice(valor_total_dolar, nome, description):    
        data_inicio = datetime.now().strftime('%d/%m/%Y')
        data_final = "Data Final"
        balance_due = valor_total_dolar
        amount_dados = valor_total_dolar
        rate_dados = valor_total_dolar
        balance_due = str(balance_due)
        amount_dados = str(amount_dados)
        rate_dados = str(rate_dados)

        #================================ Funções ============================
        def mudar_cor_celula(tabela, linha, celula, cor):
            cell_xml_element = tabela.rows[linha].cells[celula]._tc
            table_cell_proprieties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), cor)
            table_cell_proprieties.append(shade_obj)

        documento = Document()

        # ========================== Margem ================================
        secoes = documento.sections
        for secao in secoes:
            secao.top_margin = Cm(0.4)
            secao.bottom_margin = Cm(0.2)
            secao.left_margin = Cm(2)
            secao.right_margin = Cm(2)

        # ========================== Fonte =================================
        estilo = documento.styles['Normal']
        estilo.font.name = "Arial"
        estilo.font.size = Pt(10)

        # ========================= Parte superior =========================
        linha_1 = documento.add_paragraph('')
        linha_1.add_run("WIDE BRAZIL PEOPLE RECRUT. ESP. E                                                             ").bold = True
        linha_1.paragraph_format.space_after = Pt(0)
        logo = linha_1.add_run()
        logo.add_picture('logo_fundo_branco.png', width=Inches(1.5), height=Inches(0.5))
        linha_2 = documento.add_paragraph('')
        linha_2.add_run('SERV. CORP. LTDA').bold = True
        linha_2.paragraph_format.space_after = Pt(0)
        linha_4 = documento.add_paragraph('1030 Alameda Rio Negro- Escritório 206                                                       ')
        linha_4.paragraph_format.space_after = Pt(0)
        linha_5 = documento.add_paragraph('Barueri, SP 06454-000 BR')
        linha_5.paragraph_format.space_after = Pt(0)
        linha_6 = documento.add_paragraph('info@widebrazil.com')
        linha_6.paragraph_format.space_after = Pt(0)
        linha_7 = documento.add_paragraph('CNPJ TAX ID 41450051000100')
        linha_7.paragraph_format.space_after = Pt(0)
        linha_8 = documento.add_paragraph().add_run('INVOICE')
        linha_8.font.size = Pt(20)
        linha_8.font.color.rgb = RGBColor(17, 83, 7)
        linha_9 = documento.add_paragraph('')
        linha_9.add_run('BILL TO').bold = True
        linha_9.paragraph_format.space_after = Pt(0)
        linha_9.add_run('                                                                                                    INVOICE NO. ').bold = True
        linha_9.paragraph_format.space_before = Pt(2)
        linha_9.paragraph_format.space_after = Pt(0)
        linha_9.add_run(' WORKA_secdep_IM')
        linha_9.paragraph_format.space_after = Pt(0)
        linha_10 = documento.add_paragraph('WORCA PTE. LTD.')
        linha_10.paragraph_format.space_after = Pt(0)
        linha_10.add_run(f'                                                                                               DATE  ').bold = True
        linha_10.add_run(f'{data_inicio}')
        linha_10.paragraph_format.space_after = Pt(0)
        linha_11 = documento.add_paragraph('Mr. Chia-yu Wu')
        linha_11.paragraph_format.space_after = Pt(0)
        linha_11.add_run(f'                                                                                             DUE DATE  ').bold = True
        linha_11.add_run(f'{data_final}')
        linha_11.paragraph_format.space_after = Pt(0)
        linha_12 = documento.add_paragraph('68 Circular Road#02-01 -')
        linha_12.paragraph_format.space_after = Pt(0)
        linha_13 = documento.add_paragraph('Singapore (049422)')
        linha_13.paragraph_format.space_after = Pt(0)

        # ========================= Linha ========================
        linha_1 = documento.add_paragraph().add_run('_________________________________________________________________________________________')
        linha_1.font.color.rgb = RGBColor(42, 101, 28)
        #pular_linha = documento.add_paragraph()
        #pular_linha = documento.add_paragraph()

        # ========================= Meio =========================
        tabela_1 = documento.add_table(rows=1, cols=4)

        # Pintando Tabela
        mudar_cor_celula(tabela_1, 0, 0, 'd4e0d2')
        mudar_cor_celula(tabela_1, 0, 1, 'd4e0d2')
        mudar_cor_celula(tabela_1, 0, 2, 'd4e0d2')
        mudar_cor_celula(tabela_1, 0, 3, 'd4e0d2')

        # Preenchendo tabela_1
        tabela_1_linha_1 = tabela_1.rows[0].cells
        tabela_1_linha_1[0].text = ''
        tabela_1_linha_1[1].paragraphs[0].add_run('                      DESCRIPTION').bold = False
        tabela_1.rows[0].cells[1].width = Cm(50)
        tabela_1_linha_1[2].paragraphs[0].add_run('RATE').bold = False
        tabela_1.rows[0].cells[2].width = Cm(5)
        tabela_1_linha_1[3].paragraphs[0].add_run('AMOUNT').bold = False
        tabela_1.style = 'Colorful Grid Accent 1'

        # Mudando fonte tabela_1
        for row in tabela_1.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:    
                    for run in paragraph.runs:   
                        font = run.font          
                        font.size= Pt(10)  
                        font.color.rgb = RGBColor(42, 101, 28)

        # Pegando quantidades de proventos
        numero_proventos = 0
        for chave in description.keys():
            if description[chave] != "":
                numero_proventos += 1
        print(f"numero_proventos: {numero_proventos}")

        # Preenchendo tabela_2
        tabela_2 = documento.add_table(rows=numero_proventos+1, cols=3)
        tabela_2.style = "Medium List 1 Accent 3"
        tabela_2_linha = tabela_2.rows[0].cells
        tabela_2.rows[0].cells[0].width = Cm(18)
        #tabela_2_linha[0].paragraphs[0].add_run('Security Deposit').bold = True
        linha_tabela = 2
        if description["salario"] != "":        
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Salário:............................................................................R$ {description["salario"]}')
            linha_tabela += 1        
        if description["comissao"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Comissão:.......................................................................R$ {description["comissao"]}')
            linha_tabela += 1
        if description["dsr_comissao"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'D.S.R. Sobre Comissão:.................................................R$ {description["dsr_comissao"]}')
            linha_tabela += 1
        if description["vale_refeicao"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Vale Refeicao Alimentacao:............................................R$ {description["vale_refeicao"]}')
            linha_tabela += 1
        if description["reembolso_despesas"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Reembolso de despesas:................................................R$ {description["reembolso_despesas"]}')
            linha_tabela += 1
        if description["reembolso_saude"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Reembolso Seg. Saúde:.................................................R$ {description["reembolso_saude"]}')
            linha_tabela += 1
        if description["previdencia_privada"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Previdência Privada:.......................................................R$ {description["previdencia_privada"]}')
            linha_tabela += 1
        if description["base_inss_empresa"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Base INSS Empresa:......................................................R$ {description["base_inss_empresa"]}')
            linha_tabela += 1
        if description["base_fgts"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Base F.G.T.S.:................................................................R$ {description["base_fgts"]}')
            linha_tabela += 1
        if description["flash"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Flash:..............................................................................R$ {description["flash"]}')
            linha_tabela += 1
        if description["ferias"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Férias:.............................................................................R$ {description["ferias"]}')
            linha_tabela += 1
        if description["decimo_terceiro"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Décimo Terceiro:.............................................................R$ {description["decimo_terceiro"]}')
            linha_tabela += 1
        if description["seguro"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Seguro:............................................................................R$ {description["seguro"]}')
            linha_tabela += 1
        if description["adiantamento_anterior"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Adiantamento:..................................................................R$ {description["adiantamento_anterior"]}')
            linha_tabela += 1
        if description["inss_salario"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'INSS Salário:..................................................................R$ {description["inss_salario"]}')
            linha_tabela += 1
        if description["irrf_salario"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'IRRF Salário:..................................................................R$ {description["irrf_salario"]}')
            linha_tabela += 1
        if description["liquido"] != "":
            tabela_2_linha = tabela_2.rows[linha_tabela-1].cells
            tabela_2_linha[1].text = (f'Líquído:............................................................................R$ {description["liquido"]}')
            linha_tabela += 1

        tabela_2.rows[1].cells[1].width = Cm(35)
        tabela_2.columns[1].cells[1].height = Cm(5)
        tabela_2_linha = tabela_2.rows[0].cells
        tabela_2_linha[1].text = f"                                                                                                             {rate_dados}"
        tabela_2_linha[2].text = amount_dados

        # Mudando fonte tabela_2
        for row in tabela_2.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:    
                    for run in paragraph.runs:   
                        font = run.font          
                        font.size= Pt(10)

        # ========================= Linha =========================
        linha_2 = documento.add_paragraph().add_run('-----------------------------------------------------------------------------------------------------------------------------------------------------')
        linha_2.font.color.rgb = RGBColor(179, 179, 179) 

        # ========================= Baixo =========================
        baixo_linha_1 = documento.add_paragraph()
        baixo_linha_1.add_run('Payable to: WIDE BRAZIL PEOPLE RECR ESP E SERV                    ').font.size = Pt(8)
        baixo_linha_1.add_run('BALANCE DUE                      ').font.size = Pt(10)
        usd_balance_due = baixo_linha_1.add_run(f'USD {balance_due}')
        usd_balance_due.font.size = Pt(16)
        usd_balance_due.bold = True
        baixo_linha_1.paragraph_format.space_after = Pt(1)
        baixo_linha_2 = documento.add_paragraph().add_run('CORPORATIVOS LTDA \n'\
                                                        'Intermediary Institution (or Correspondent Bank, Field 56)\n'\
                                                        'SWIFT Code: CITIUS33 \n'\
                                                        'Bank Name: CITI, NEW YORK - Bank Country: UNITED STATES \n'\
                                                        'Beneficiary Bank (Field 57):\n'\
                                                        'SWIFT Code: BPABBRRJ\n'\
                                                        'BTG Account with CITI, NEW YORK: 36317173\n'\
                                                        'Bank Name: BANCO BTG PACTUAL S.A.\n'\
                                                        'Bank Address: PRAIA DE BOTAFOGO, 501, RIO DE JANEIRO\n'\
                                                        'Bank Country: BRAZIL \n'\
                                                        'Beneficiary Customer (Field 59):\n'\
                                                        'Beneficiary IBAN: BR2930306294000010002827654C1\n'\
                                                        'Beneficiary Name: WIDE BRAZIL PEOPLE RECRUTAMENTO\n'\
                                                        'ESPECIALIZADO E SERVICOS CORPORATIVOS LTDA\n'\
                                                        'Beneficiary Address: Commercial - AL RIO NEGRO, 1030 -\n'\
                                                        'ALPHAVILLE CENTRO -\n'\
                                                        'BARUERÍ - SÃO PAULO - 06454000\n'\
                                                        'Beneficiary Country: BRAZIL\n'\
                                                        'Remittance Information (Field 70):\n'\
                                                        'If possible, include invoice number on field 70\n'  
                                                        )
        baixo_linha_2.font.size = Pt(8)

        # =============================== RODAPE ==================================
        rodape = documento.add_paragraph().add_run('\n                                                                      Please let us know if you have any issues with the payment.\n\n'\
                                        '                                                                             As always, thank you very much for your business.'
                                        )
        rodape.font.size = Pt(8)

        documento.save(f"invoice.docx")
        #convert(f"invoice.docx", "invoice.pdf")
    def pegar_valor_conta_receber(codigo_cliente_omie, description):
        pagina = 1
        total_de_paginas = 1
        valor_total = 0
        while pagina <= total_de_paginas:
            url = "https://app.omie.com.br/api/v1/financas/contareceber/"
            payload = json.dumps({
                                    "call": "ListarContasReceber",
                                    "app_key": app_key,
                                    "app_secret": app_secret,
                                    "param":[
                                                {
                                                    "pagina": pagina,
                                                    "registros_por_pagina": 500,
                                                    "apenas_importado_api": "N"
                                                }
                                            ]
                                })
            headers ={
                        'Content-Type': 'application/json'
                    }
            response = requests.request("POST", url, headers=headers, data=payload)
            response = response.json()
            pagina = response["pagina"]
            total_de_paginas = response["total_de_paginas"]
            conta_receber_cadastro = response["conta_receber_cadastro"]
            for conta_receber in conta_receber_cadastro:
                codigo_cliente_fornecedor = conta_receber["codigo_cliente_fornecedor"]
                status_titulo = conta_receber["status_titulo"]
                if codigo_cliente_fornecedor == codigo_cliente_omie and status_titulo != "RECEBIDO":
                    valor_total += float(conta_receber["valor_documento"])
                    categorias = conta_receber["categorias"]
                    categorias = categorias[0]
                    codigo_categoria = categorias["codigo_categoria"]
                    valor_documento = conta_receber["valor_documento"]
                    valor_documento = (f'{valor_documento:,.2f}')
                    valor_documento = valor_documento.replace(",", "_")
                    valor_documento = valor_documento.replace(".", ",")
                    valor_documento = valor_documento.replace("_", ".")
                    if codigo_categoria == categoria_receber_salario:
                        description["salario"] = valor_documento
                    if codigo_categoria == categoria_receber_comissao:
                        description["comissao"] = valor_documento
                    if codigo_categoria == categoria_receber_dsr:
                        description["dsr_comissao"] = valor_documento
                    if codigo_categoria == categoria_receber_alimentacao:
                        description["vale_refeicao"] = valor_documento
                    if codigo_categoria == categoria_receber_reembolso_despesas:
                        description["reembolso_despesas"] = valor_documento
                    if codigo_categoria == categoria_receber_reembolso_saude:
                        description["reembolso_saude"] = valor_documento
                    if codigo_categoria == categoria_receber_previdencia:
                        description["previdencia_privada"] = valor_documento
                    if codigo_categoria == categoria_receber_inss_empresa:
                        description["base_inss_empresa"] = valor_documento
                    if codigo_categoria == categoria_receber_fgts:
                        description["base_fgts"] = valor_documento
                    if codigo_categoria == categoria_receber_flash:
                        description["flash"] = valor_documento
                    if codigo_categoria == categoria_receber_ferias:
                        description["ferias"] = valor_documento
                    if codigo_categoria == categoria_receber_decimo:
                        description["decimo_terceiro"] = valor_documento
                    if codigo_categoria == categoria_receber_seguro:
                        description["seguro"] = valor_documento
                    if codigo_categoria == categoria_receber_liquido:
                        description["liquido"] = valor_documento
                    if codigo_categoria == categoria_receber_adiantamento:
                        description["adiantamento_anterior"] = valor_documento
            pagina += 1
        return valor_total, description
    def pegar_data_vencimento():
        data_atual = date.today()
        data_vencimento = data_atual + timedelta(days=30)
        data_atual = data_atual.strftime("%d/%m/%Y")
        data_vencimento = data_vencimento.strftime("%d/%m/%Y")
        return data_vencimento
    def incluir_conta_receber_codigo(codigo_cliente_omie, data_vencimento, valor_documento, codigo_categoria):
        randomlist = random.sample(range(1, 12), 8)
        randomlist = str(randomlist)
        aleatorio = randomlist.replace(",","")
        aleatorio = aleatorio.replace(" ","")
        aleatorio = aleatorio.replace("[","")
        codigo_lancamento_integracao = aleatorio.replace("]","")    
        url = "https://app.omie.com.br/api/v1/financas/contareceber/"
        payload = json.dumps({
                                "call": "IncluirContaReceber",
                                "app_key": app_key,
                                "app_secret": app_secret,
                                "param":[
                                            {
                                                "codigo_lancamento_integracao": codigo_lancamento_integracao,
                                                "codigo_cliente_fornecedor": codigo_cliente_omie,
                                                "data_vencimento": data_vencimento,
                                                "valor_documento": valor_documento,
                                                "codigo_categoria": codigo_categoria,
                                                "id_conta_corrente": "7311700205"
                                            }
                                        ]
                            })
        headers ={
                    'Content-Type': 'application/json'
                }
        response = requests.request("POST", url, headers=headers, data=payload)
        response = response.json()
        print(f'IncluirContaReceber: {response}')
        codigo_lancamento_omie = response["codigo_lancamento_omie"]
        return codigo_lancamento_omie
    def anexar_invoice(nId, cArquivo):
        cMd5 = "123"
        while True:      
            url = "https://app.omie.com.br/api/v1/geral/anexo/"
            payload = json.dumps({
                                    "call": "IncluirAnexo",
                                    "app_key": app_key,
                                    "app_secret": app_secret,
                                    "param":[
                                                {
                                                    "cTabela": "conta-receber",
                                                    "nId": nId,
                                                    "cNomeArquivo": "invoice.docx",
                                                    "cTipoArquivo": "",
                                                    "cArquivo": cArquivo,
                                                    "cMd5": cMd5
                                                }
                                            ]
                                })
            headers ={
                        'Content-Type': 'application/json'
                    }
            response = requests.request("POST", url, headers=headers, data=payload)
            response = response.json()
            print(f'IncluirAnexo: {response}')
            try:
                faultstring = response["faultstring"]
                faultstring = faultstring.split(" ")
                faultstring = faultstring[9]
                faultstring = str(faultstring)
                faultstring = faultstring.replace("[", "")
                faultstring = faultstring.replace("]", "")
                cMd5 = faultstring.replace("!", "")
                print(f"cMd5: {cMd5}")      
            except:
                break

    #====================== Recebendo Arquivo S3 ======================#
    s3 = boto3.resource("s3", aws_access_key_id="AKIATX77KZ6NA7RTXMFO", aws_secret_access_key="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")
    bucket = s3.Bucket("parceiro-do-contador-bucket")
    bucket.download_file(Key="arquivos/j_son.json", Filename="j_son.json")
    bucket.download_file(Key="arquivos/logo_fundo_branco.png", Filename="logo_fundo_branco.png")
    with open('j_son.json', 'r') as arquivo:
        j_son = arquivo.read()
    j_son = j_son.replace("'", "\"")
    j_son = json.loads(j_son)
    nome = j_son["nome"]
    print(f'nome: {nome}')
    try:
        nome = nome.encode("windows-1252").decode("utf-8")
    except:
        pass
    cotacao_dolar = j_son["cotacao_dolar"]
    if "," in str(cotacao_dolar):
        cotacao_dolar = cotacao_dolar.replace(",", ".")
    cotacao_dolar = float(cotacao_dolar)

    #codigo_cliente_omie = buscar_codigo_cliente_teste(nome)
    codigo_cliente_omie = buscar_codigo_cliente(nome)
    description = {"salario": "", "comissao": "", "dsr_comissao": "",  "vale_refeicao": "", "reembolso_despesas": "", "reembolso_saude": "", "inss_salario": "",\
    "adiantamento_anterior": "", "irrf_salario": "", "previdencia_privada": "", "base_inss_empresa": "", "base_fgts": "", "liquido": "", "ferias": "",\
    "decimo_terceiro": "", "seguro": "", "flash": ""}            
    valor_total, description = pegar_valor_conta_receber(codigo_cliente_omie, description)                
    valor_total_dolar = valor_total / cotacao_dolar
    valor_total_dolar = (f'{valor_total_dolar:,.2f}')
    valor_total_dolar = valor_total_dolar.replace(",", "_")
    valor_total_dolar = valor_total_dolar.replace(".", ",")
    valor_total_dolar = valor_total_dolar.replace("_", ".")
    gerar_invoice(valor_total_dolar, nome, description)
    shutil.make_archive('invoice', 'zip', '', 'invoice.docx')
    with open("invoice.zip", "rb") as arquivo:
        invoice_64 = base64.b64encode(arquivo.read())
    invoice_64 = str(invoice_64)
    invoice_64 = invoice_64[:-1]
    invoice_64 = invoice_64[2:]
    codigo_cliente_omie = buscar_codigo_cliente(nome)
    data_vencimento = pegar_data_vencimento()
    codigo_lancamento_omie = incluir_conta_receber_codigo(codigo_cliente_omie, data_vencimento, valor_total, categoria_invoice)
    anexar_invoice(codigo_lancamento_omie, invoice_64)